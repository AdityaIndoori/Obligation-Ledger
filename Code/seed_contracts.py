"""Generate the four demo contracts from MASTER doc T13, plus their replay
fixtures, so the whole pipeline runs before any model is serving.

Deliberately spread across formats to prove the ingest matrix on stage:
  meridian-msa    .docx  the star -- renewal buried in 14.3, fee table
  delta-sow       .pdf   payment milestones, no auto-renewal
  northgate-nda   .txt   short and clean, processes fast
  acme-services   .doc   HELD BACK for the offline moment (legacy format,
                         so the unplug beat also demonstrates Tier 2)

Fixtures record what a correct extraction looks like for each. They are
hand-authored ONLY so the app can be built and verified now; they are replaced
by real recordings (LLM_MODE=record) once the model lane is up. Each is marked
as hand-authored so nothing can mistake one for a model recording.
"""
import hashlib
import json
import os
import subprocess
import sys
import tempfile
from datetime import date, timedelta

OUT = os.environ.get("SEED_OUT", "/srv/ledger/intake")
HOLD = os.path.join(OUT, "..", "holdback")
FIXTURES = os.environ.get("LEDGER_FIXTURES", "/srv/ledger/data/fixtures")

# --------------------------------------------------------------- MERIDIAN
MERIDIAN_HEAD = [
    "MASTER SERVICES AGREEMENT",
    "",
    "This Master Services Agreement (this \"Agreement\") is entered into by and "
    "between Meridian Holdings LLC, a Delaware limited liability company "
    "(\"Client\"), and Northgate Legal Services LLP, a limited liability "
    "partnership (\"Provider\").",
    "",
    "1. TERM. This Agreement is effective January 1, 2026 and shall remain in "
    "force until March 31, 2027, unless earlier terminated in accordance with "
    "Section 12.",
    "",
    "2. SERVICES. Provider shall furnish the professional services described in "
    "each Statement of Work executed by the parties.",
    "",
    "3. FEES. Client shall pay Provider USD 120,000 per annum, invoiced "
    "quarterly in arrears, due within thirty (30) days of invoice date.",
    "",
    "4. EXPENSES. Reasonable out-of-pocket expenses are reimbursable when "
    "pre-approved in writing.",
    "",
    "5. CONFIDENTIALITY. Each party shall protect the other's Confidential "
    "Information using no less than reasonable care.",
    "",
    "6. INTELLECTUAL PROPERTY. Work product prepared under a Statement of Work "
    "vests in Client upon payment in full.",
    "",
    "7. DATA PROTECTION. Provider shall process Client personal data solely on "
    "documented instructions from Client.",
    "",
    "8. WARRANTIES. Provider warrants that services will be performed in a "
    "professional and workmanlike manner.",
    "",
    "9. LIMITATION OF LIABILITY. Except as stated in Section 11, neither "
    "party's aggregate liability shall exceed the fees paid in the twelve (12) "
    "months preceding the claim.",
    "",
    "10. INSURANCE. Provider shall maintain professional indemnity cover of not "
    "less than USD 5,000,000 per occurrence.",
    "",
    "11. INDEMNITY. Provider shall indemnify Client against third-party claims "
    "arising from Provider's negligence. Provider's indemnity obligations "
    "under Section 11.2 shall be uncapped.",
    "",
    "12. TERMINATION. Either party may terminate for material breach not cured "
    "within thirty (30) days of written notice.",
    "",
    "13. FORCE MAJEURE. Neither party is liable for delay caused by events "
    "beyond its reasonable control.",
    "",
    "14. GENERAL.",
    "",
    "14.1 Assignment. Neither party may assign this Agreement without the "
    "other's prior written consent.",
    "",
    "14.2 Notices. Notices must be given in writing to the addresses set out "
    "in the signature block.",
    "",
    # ---- the clause the whole demo turns on, buried late, as specified.
    "14.3 Renewal. This Agreement shall automatically renew for successive "
    "twelve-month terms unless either party gives written notice at least "
    "sixty (60) days prior to expiry.",
    "",
    "14.4 Governing Law. This Agreement is governed by the laws of the State "
    "of Delaware, without regard to its conflict-of-laws rules.",
    "",
    "14.5 Entire Agreement. This Agreement constitutes the entire agreement "
    "between the parties with respect to its subject matter.",
]

MERIDIAN_RENEWAL_SPAN = (
    "This Agreement shall automatically renew for successive twelve-month "
    "terms unless either party gives written notice at least sixty (60) days "
    "prior to expiry.")

MERIDIAN_DATA = {
    "parties": [
        {"name": "Meridian Holdings LLC", "role": "Client",
         "source_span": "Meridian Holdings LLC, a Delaware limited liability company"},
        {"name": "Northgate Legal Services LLP", "role": "Provider",
         "source_span": "Northgate Legal Services LLP, a limited liability partnership"},
    ],
    "effective_date": {"value": "2026-01-01",
                       "source_span": "This Agreement is effective January 1, 2026"},
    "term_end": {"value": "2027-03-31",
                 "source_span": "shall remain in force until March 31, 2027"},
    "auto_renewal": {"present": True, "renewal_term_months": 12,
                     "notice_days": 60,
                     "source_span": MERIDIAN_RENEWAL_SPAN},
    "payment": {"amount": "120,000", "currency": "USD", "schedule": "quarterly in arrears",
                "source_span": "Client shall pay Provider USD 120,000 per annum, invoiced quarterly in arrears"},
    "governing_law": {"value": "Delaware",
                      "source_span": "governed by the laws of the State of Delaware"},
    "unusual_terms": [
        {"summary": "Provider indemnity under 11.2 is uncapped",
         "why_unusual": "Uncapped indemnity sits outside the Section 9 liability cap",
         "source_span": "Provider's indemnity obligations under Section 11.2 shall be uncapped."},
    ],
}

# --------------------------------------------------------------- DELTA
DELTA_TEXT = "\n\n".join([
    "STATEMENT OF WORK - DELTA CONSULTING GROUP",
    "This Statement of Work is entered into between Delta Consulting Group Inc. "
    "(\"Consultant\") and Meridian Holdings LLC (\"Client\").",
    "1. COMMENCEMENT. This Statement of Work is effective February 15, 2026.",
    "2. COMPLETION. All work shall be complete by December 31, 2026. This "
    "Statement of Work expires on that date and does not renew.",
    "3. FEES AND MILESTONES. Total fees are USD 85,000, payable against "
    "milestones: USD 25,000 on commencement; USD 30,000 on delivery of the "
    "interim report on October 30, 2026; USD 30,000 on acceptance.",
    "4. ACCEPTANCE. Client shall have ten (10) business days to accept or "
    "reject each deliverable.",
    "5. PERSONNEL. Consultant shall not substitute named personnel without "
    "Client's written consent.",
    "6. INDEPENDENT CONTRACTOR. Consultant is an independent contractor and "
    "not an employee of Client.",
    "7. GOVERNING LAW. This Statement of Work is governed by the laws of the "
    "State of New York.",
    "8. NO AUTOMATIC RENEWAL. The parties expressly agree that this Statement "
    "of Work contains no automatic renewal provision. Any extension requires a "
    "written amendment signed by both parties.",
])

DELTA_DATA = {
    "parties": [
        {"name": "Delta Consulting Group Inc.", "role": "Consultant",
         "source_span": "Delta Consulting Group Inc. (\"Consultant\")"},
        {"name": "Meridian Holdings LLC", "role": "Client",
         "source_span": "Meridian Holdings LLC (\"Client\")"},
    ],
    "effective_date": {"value": "2026-02-15",
                       "source_span": "This Statement of Work is effective February 15, 2026."},
    "term_end": {"value": "2026-12-31",
                 "source_span": "All work shall be complete by December 31, 2026."},
    "auto_renewal": {"present": False, "renewal_term_months": None,
                     "notice_days": None,
                     "source_span": "this Statement of Work contains no automatic renewal provision"},
    "payment": {"amount": "85,000", "currency": "USD", "schedule": "milestones",
                "source_span": "Total fees are USD 85,000, payable against "
                               "milestones: USD 25,000 on commencement; USD "
                               "30,000 on delivery of the interim report on "
                               "October 30, 2026"},
    "governing_law": {"value": "New York",
                      "source_span": "governed by the laws of the State of New York"},
    "unusual_terms": [],
}

# --------------------------------------------------------------- NORTHGATE
NORTHGATE_TEXT = "\n\n".join([
    "MUTUAL NON-DISCLOSURE AGREEMENT",
    "This Mutual Non-Disclosure Agreement is made between Northgate Legal "
    "Services LLP and Acme Industrial Services Inc.",
    "1. EFFECTIVE DATE. This Agreement is effective March 1, 2026.",
    "2. TERM. This Agreement expires on March 1, 2028.",
    "3. PURPOSE. The parties wish to exchange confidential information to "
    "evaluate a potential commercial relationship.",
    "4. CONFIDENTIALITY. Each party shall keep the other's Confidential "
    "Information confidential and use it solely for the Purpose.",
    "5. EXCLUSIONS. Information that is public, independently developed, or "
    "lawfully received from a third party is excluded.",
    "6. RETURN. Upon request each party shall return or destroy the other's "
    "Confidential Information.",
    "7. NO LICENCE. Nothing in this Agreement grants any licence to "
    "intellectual property.",
    "8. GOVERNING LAW. This Agreement is governed by the laws of England and "
    "Wales.",
])

NORTHGATE_DATA = {
    "parties": [
        {"name": "Northgate Legal Services LLP", "role": "Discloser",
         "source_span": "Northgate Legal Services LLP"},
        {"name": "Acme Industrial Services Inc.", "role": "Recipient",
         "source_span": "Acme Industrial Services Inc."},
    ],
    "effective_date": {"value": "2026-03-01",
                       "source_span": "This Agreement is effective March 1, 2026."},
    "term_end": {"value": "2028-03-01",
                 "source_span": "This Agreement expires on March 1, 2028."},
    "auto_renewal": {"present": False, "renewal_term_months": None,
                     "notice_days": None, "source_span": ""},
    "payment": {"amount": "", "currency": "", "schedule": "", "source_span": ""},
    "governing_law": {"value": "England and Wales",
                      "source_span": "governed by the laws of England and Wales"},
    "unusual_terms": [],
}

# --------------------------------------------------------------- ACME (held back)
ACME_TEXT = "\n\n".join([
    "SERVICES AGREEMENT - ACME INDUSTRIAL SERVICES",
    "This Services Agreement is made between Acme Industrial Services Inc. "
    "(\"Supplier\") and Meridian Holdings LLC (\"Customer\").",
    "1. TERM. This Agreement is effective April 1, 2026 and continues until "
    "September 30, 2027.",
    "2. CHARGES. Customer shall pay USD 240,000 per annum, invoiced monthly.",
    "3. ESCALATION. Charges increase on each anniversary by CPI plus three "
    "percent (3%).",
    "4. SERVICE LEVELS. Supplier shall meet the availability targets in "
    "Schedule 2.",
    "5. AUDIT. Customer may audit Supplier's compliance once per contract year.",
    "6. SUBCONTRACTING. Supplier shall not subcontract without Customer's "
    "written consent.",
    "7. RENEWAL. This Agreement renews automatically for further periods of "
    "twenty-four (24) months unless either party serves written notice not "
    "less than ninety (90) days before the end of the then-current term.",
    "8. GOVERNING LAW. This Agreement is governed by the laws of the State of "
    "Illinois.",
])

ACME_RENEWAL_SPAN = (
    "This Agreement renews automatically for further periods of twenty-four "
    "(24) months unless either party serves written notice not less than "
    "ninety (90) days before the end of the then-current term.")

ACME_DATA = {
    "parties": [
        {"name": "Acme Industrial Services Inc.", "role": "Supplier",
         "source_span": "Acme Industrial Services Inc. (\"Supplier\")"},
        {"name": "Meridian Holdings LLC", "role": "Customer",
         "source_span": "Meridian Holdings LLC (\"Customer\")"},
    ],
    "effective_date": {"value": "2026-04-01",
                       "source_span": "This Agreement is effective April 1, 2026"},
    "term_end": {"value": "2027-09-30",
                 "source_span": "continues until September 30, 2027."},
    "auto_renewal": {"present": True, "renewal_term_months": 24,
                     "notice_days": 90, "source_span": ACME_RENEWAL_SPAN},
    "payment": {"amount": "240,000", "currency": "USD", "schedule": "monthly",
                "source_span": "Customer shall pay USD 240,000 per annum, invoiced monthly."},
    "governing_law": {"value": "Illinois",
                      "source_span": "governed by the laws of the State of Illinois"},
    "unusual_terms": [
        {"summary": "Charges escalate annually by CPI plus 3%",
         "why_unusual": "Uncapped index-linked escalator compounds each year",
         "source_span": "Charges increase on each anniversary by CPI plus three percent (3%)."},
    ],
}

# ---------------------------------------------------- STERLING (near-term)
# Deliberately dated so its notice deadline falls a few weeks out, which is
# what makes the Deadlines board and the Ask tab show live urgency on stage
# instead of dates two years away.
_STERLING_END = (date.today() + timedelta(days=105))
_STERLING_EFF = (_STERLING_END - timedelta(days=730))


def _long(d):
    return d.strftime("%B ") + str(d.day) + d.strftime(", %Y")


STERLING_TEXT = "\n\n".join([
    "SOFTWARE LICENCE AND SUPPORT AGREEMENT",
    "This Agreement is made between Sterling Systems Ltd (\"Licensor\") and "
    "Meridian Holdings LLC (\"Licensee\").",
    f"1. TERM. This Agreement is effective {_long(_STERLING_EFF)} and shall "
    f"continue until {_long(_STERLING_END)}.",
    "2. LICENCE. Licensor grants Licensee a non-exclusive licence to use the "
    "Software for internal business purposes.",
    "3. SUPPORT. Licensor shall provide support during business hours with a "
    "four (4) hour response target for critical incidents.",
    "4. FEES. Licensee shall pay USD 96,000 per annum, invoiced annually in "
    "advance.",
    "5. AUDIT RIGHTS. Licensor may audit Licensee's use of the Software on "
    "thirty (30) days notice, not more than once per year.",
    "6. DATA. Licensee data remains the property of Licensee at all times.",
    "7. RENEWAL. This Agreement shall renew automatically for further terms of "
    "twelve (12) months unless either party gives written notice at least "
    "thirty (30) days prior to the expiry of the then-current term.",
    "8. LIMITATION. Licensor's liability is limited to the fees paid in the "
    "preceding twelve months, except that Licensor's liability for breach of "
    "confidentiality shall be unlimited.",
    "9. GOVERNING LAW. This Agreement is governed by the laws of England and "
    "Wales.",
])

STERLING_RENEWAL_SPAN = (
    "This Agreement shall renew automatically for further terms of twelve (12) "
    "months unless either party gives written notice at least thirty (30) days "
    "prior to the expiry of the then-current term.")

STERLING_DATA = {
    "parties": [
        {"name": "Sterling Systems Ltd", "role": "Licensor",
         "source_span": "Sterling Systems Ltd (\"Licensor\")"},
        {"name": "Meridian Holdings LLC", "role": "Licensee",
         "source_span": "Meridian Holdings LLC (\"Licensee\")"},
    ],
    "effective_date": {"value": _STERLING_EFF.isoformat(),
                       "source_span": f"This Agreement is effective {_long(_STERLING_EFF)}"},
    "term_end": {"value": _STERLING_END.isoformat(),
                 "source_span": f"shall continue until {_long(_STERLING_END)}."},
    "auto_renewal": {"present": True, "renewal_term_months": 12,
                     "notice_days": 30, "source_span": STERLING_RENEWAL_SPAN},
    "payment": {"amount": "96,000", "currency": "USD", "schedule": "annually in advance",
                "source_span": "Licensee shall pay USD 96,000 per annum, invoiced annually in advance."},
    "governing_law": {"value": "England and Wales",
                      "source_span": "governed by the laws of England and Wales"},
    "unusual_terms": [
        {"summary": "Unlimited liability for breach of confidentiality",
         "why_unusual": "Carve-out escapes the twelve-month fee cap entirely",
         "source_span": "Licensor's liability for breach of confidentiality shall be unlimited."},
    ],
}


def write_meridian_docx(path):
    """Body paragraphs + a fee table + a header. The table matters: it is where
    a real fee schedule lives, and a paragraph-only parser would drop it."""
    from docx import Document
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = \
        "CONFIDENTIAL - Meridian Holdings LLC / Northgate Legal Services LLP"
    for para in MERIDIAN_HEAD:
        doc.add_paragraph(para)
    doc.add_paragraph("SCHEDULE 1 - FEE SCHEDULE")
    table = doc.add_table(rows=4, cols=2)
    for i, (k, v) in enumerate([
            ("Annual fee", "USD 120,000"),
            ("Invoicing", "Quarterly in arrears"),
            ("Payment terms", "Thirty (30) days from invoice date"),
            ("Renewal term", "Twelve (12) months")]):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
    doc.save(path)


def write_pdf(path, text, title):
    import fitz
    doc = fitz.open()
    words, page_chunks, cur = text.split(), [], []
    # ~1600 chars per page keeps the SOW to a realistic multi-page document
    size = 0
    for w in words:
        cur.append(w)
        size += len(w) + 1
        if size > 1600:
            page_chunks.append(" ".join(cur))
            cur, size = [], 0
    if cur:
        page_chunks.append(" ".join(cur))
    for chunk in page_chunks:
        page = doc.new_page()
        page.insert_textbox(fitz.Rect(56, 56, 540, 760), chunk, fontsize=10.5,
                            fontname="helv")
    doc.save(path)


def write_legacy_doc(path, text):
    """Write a real legacy .doc via LibreOffice, so the held-back contract also
    exercises the Tier 2 conversion path on stage."""
    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, os.path.basename(path).replace(".doc", ".txt"))
        with open(src, "w") as fh:
            fh.write(text)
        subprocess.run(["soffice", "--headless", "--norestore", "--convert-to",
                        "doc", "--outdir", tmp, src],
                       check=True, timeout=180,
                       stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        produced = src[:-4] + ".doc"
        if not os.path.exists(produced):
            raise RuntimeError("LibreOffice did not produce a .doc")
        os.replace(produced, path)


def write_odt(path, text):
    from odf.opendocument import OpenDocumentText
    from odf.text import P
    doc = OpenDocumentText()
    for para in text.split("\n\n"):
        doc.text.addElement(P(text=para))
    doc.save(path)


def fixture_for(text, data, note):
    os.makedirs(FIXTURES, exist_ok=True)
    sha = hashlib.sha256(text.encode("utf-8", "replace")).hexdigest()
    dest = os.path.join(FIXTURES, f"{sha}.json")
    with open(dest, "w") as fh:
        json.dump({"model": "hand-authored-fixture", "data": data,
                   "raw": "(hand-authored: replace with LLM_MODE=record)",
                   "hand_authored": True, "note": note}, fh, indent=1)
    return dest


def main():
    os.makedirs(OUT, exist_ok=True)
    os.makedirs(HOLD, exist_ok=True)
    import ingest

    made = []

    # 1. Meridian -- the star, as .docx
    mer = os.path.join(OUT, "meridian-msa.docx")
    write_meridian_docx(mer)
    made.append(mer)

    # 2. Delta -- as a real multi-page PDF
    delta = os.path.join(OUT, "delta-sow.pdf")
    write_pdf(delta, DELTA_TEXT, "Delta SOW")
    made.append(delta)

    # 3. Northgate -- plain text, fast
    ng = os.path.join(OUT, "northgate-nda.txt")
    with open(ng, "w") as fh:
        fh.write(NORTHGATE_TEXT)
    made.append(ng)

    # 3b. Sterling -- near-term notice deadline, as .odt (widens the format
    #     spread and gives the Deadlines board real urgency on stage)
    sterling = os.path.join(OUT, "sterling-licence.odt")
    write_odt(sterling, STERLING_TEXT)
    made.append(sterling)

    # 4. Acme -- HELD BACK, legacy .doc
    acme = os.path.join(HOLD, "acme-services.doc")
    try:
        write_legacy_doc(acme, ACME_TEXT)
    except Exception as exc:                     # noqa: BLE001
        print(f"  ! legacy .doc unavailable ({exc}); falling back to .rtf")
        acme = os.path.join(HOLD, "acme-services.rtf")
        with open(acme, "w") as fh:
            fh.write(r"{\rtf1\ansi\deff0 " +
                     ACME_TEXT.replace("\n", r"\par ") + "}")
    made.append(acme)

    # Fixtures keyed on the PARSED text, which is what extract() hashes.
    print()
    for path, data, label in ((mer, MERIDIAN_DATA, "meridian"),
                              (delta, DELTA_DATA, "delta"),
                              (ng, NORTHGATE_DATA, "northgate"),
                              (sterling, STERLING_DATA, "sterling"),
                              (acme, ACME_DATA, "acme")):
        doc = ingest.parse(path)
        dest = fixture_for(doc.text, data, f"seed:{label}")
        print(f"  {os.path.basename(path):<26} fmt={doc.fmt:<5} "
              f"chars={len(doc.text):<6} fixture={os.path.basename(dest)[:16]}...")

    print(f"\nseeded {len(made)} contracts "
          f"({len(made) - 1} in intake, 1 held back in {HOLD})")
    return 0


if __name__ == "__main__":
    sys.exit(main())
